"""Word document export for IR 1-Pager memos."""
import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


NAVY = RGBColor(0x00, 0x33, 0x66)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shading_elm)


def _add_section_header(doc, text: str):
    """Add a styled section header."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    # Add bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "003366",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_runs_with_formatting(paragraph, text: str, base_size=Pt(11), base_color=DARK_GRAY):
    """Parse inline markdown (**bold** and *italic*) and add formatted runs to a paragraph."""
    # Split on **bold** and *italic* patterns
    # Pattern: **bold text**, *italic text*, or plain text
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True
            run.font.size = base_size
            run.font.name = "Calibri"
            run.font.color.rgb = base_color
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.italic = True
            run.font.size = base_size
            run.font.name = "Calibri"
            run.font.color.rgb = base_color
        else:
            run = paragraph.add_run(part)
            run.font.size = base_size
            run.font.name = "Calibri"
            run.font.color.rgb = base_color


def _add_sub_header(doc, text: str):
    """Add a ### subsection header with subtle bottom border."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    # Add thin bottom border for visual distinction
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "2",
        qn("w:space"): "1",
        qn("w:color"): "99AABB",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_blockquote(doc, text: str):
    """Add an indented blockquote paragraph with italic styling."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    # Strip leading > and whitespace, strip wrapping quotes
    content = text.lstrip("> ").strip()
    if (content.startswith('"') and content.endswith('"')) or \
       (content.startswith('\u201c') and content.endswith('\u201d')):
        content = content[1:-1]
    run = p.add_run(f'\u201c{content}\u201d')
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = MEDIUM_GRAY
    p.space_after = Pt(4)


def _add_bullet(doc, text: str):
    """Add a bullet point with inline formatting support."""
    p = doc.add_paragraph(style="List Bullet")
    _add_runs_with_formatting(p, text)
    p.space_after = Pt(2)


def _add_numbered_item(doc, text: str):
    """Add a numbered list item with inline formatting support."""
    p = doc.add_paragraph(style="List Number")
    _add_runs_with_formatting(p, text)
    p.space_after = Pt(2)


def _add_body_text(doc, text: str):
    """Add body text paragraph with inline formatting support."""
    p = doc.add_paragraph()
    _add_runs_with_formatting(p, text)
    p.space_after = Pt(4)


def _add_markdown_table(doc, table_lines: list[str]):
    """Convert markdown table lines into a proper Word table."""
    # Parse rows, skip separator lines (|---|)
    parsed_rows = []
    for line in table_lines:
        stripped = line.strip().strip("|")
        if re.match(r'^[\s\-:|]+$', stripped):
            continue  # separator row
        cells = [c.strip() for c in stripped.split("|")]
        if cells:
            parsed_rows.append(cells)

    if not parsed_rows:
        return

    n_cols = max(len(r) for r in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(parsed_rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= n_cols:
                break
            cell = table.rows[r_idx].cells[c_idx]
            # Clear default paragraph
            cell.text = ""
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            # Strip bold markers and apply formatting
            clean = cell_text.strip()
            _add_runs_with_formatting(p, clean, base_size=Pt(10))

            # Style header row
            if r_idx == 0:
                _set_cell_shading(cell, "003366")
                for run in p.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True


def _parse_markdown_to_doc(doc, markdown_text: str):
    """Parse markdown text and add to document as structured Word content.

    Handles: ## and ### headings, **bold**, *italic*, > blockquotes,
    - bullets, 1. numbered lists, | markdown tables |, and mixed inline formatting.
    """
    # Fix common artifacts: double colons from "**response::**" pattern
    text = markdown_text.replace("::", ":")

    lines = text.strip().split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip top-level ## headers (we add our own section headers)
        if stripped.startswith("## ") and not stripped.startswith("### "):
            i += 1
            continue

        # ### Sub-headers → styled sub-heading
        if stripped.startswith("### "):
            _add_sub_header(doc, stripped[4:].strip())
            i += 1
            continue

        # Markdown table: accumulate consecutive | lines
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)
            continue

        # Blockquote (> text)
        if stripped.startswith("> "):
            _add_blockquote(doc, stripped)
            i += 1
            continue

        # Bullet point (- or *)
        if re.match(r'^[\s]*[-*]\s', stripped):
            content = re.sub(r'^[\s]*[-*]\s+', '', stripped)
            _add_bullet(doc, content)
            i += 1
            continue

        # Numbered list (1. text, 2. text, etc.)
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            _add_numbered_item(doc, num_match.group(2))
            i += 1
            continue

        # Italic line (*text* at start and end, not bold)
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run.font.color.rgb = MEDIUM_GRAY
            p.space_after = Pt(4)
            i += 1
            continue

        # Any other text — render with inline formatting
        _add_body_text(doc, stripped)
        i += 1


def _add_forecast_table(doc, forecast_data: dict):
    """Add the forecast comparison table."""
    rows = forecast_data.get("table_rows", [])
    if not rows:
        _add_body_text(doc, "[Forecast data unavailable]")
        return

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    headers = ["3Q25 Revenues ($MM)", "Analyst", "Consensus", "Δ vs Consensus %"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        _set_cell_shading(cell, "003366")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for r_idx, row in enumerate(rows):
        indent = "  " * row.get("indent", 0)
        label = f"{indent}{row['label']}"
        analyst_val = row.get("analyst")
        consensus_val = row.get("consensus")
        delta_val = row.get("delta", "")

        # Format values
        if row["label"] in ("EPS",):
            a_str = f"${analyst_val:.2f}" if analyst_val is not None else "[N/A]"
            c_str = f"${consensus_val:.2f}" if consensus_val is not None else "[N/A]"
        elif row["label"] in ("ROE", "ROTCE"):
            a_str = f"{analyst_val:.1f}%" if analyst_val is not None else "[N/A]"
            c_str = f"{consensus_val:.1f}%" if consensus_val is not None else "[N/A]"
        else:
            a_str = f"{int(analyst_val):,}" if analyst_val is not None else "[N/A]"
            c_str = f"{int(consensus_val):,}" if consensus_val is not None else "[N/A]"

        values = [label, a_str, c_str, str(delta_val)]
        for c_idx, val in enumerate(values):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    if row.get("indent", 0) == 0:
                        run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT

    # Rating and Price Target
    rating = forecast_data.get("rating", "[N/A]")
    pt = forecast_data.get("price_target")
    pt_str = f"${pt}" if pt is not None else "[N/A]"
    p = doc.add_paragraph()
    run = p.add_run(f"Rating: {rating} | Price Target: {pt_str}")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.bold = True
    p.space_before = Pt(6)


def _add_valuation_table(doc, valuation_data: dict):
    """Add the valuation ratios comparison table."""
    tickers_data = valuation_data.get("tickers", {})
    peer_median = valuation_data.get("peer_median", {})
    ticker_order = ["MS", "GS", "JPM", "BAC", "C", "WFC", "SCHW"]
    metrics = [
        "Stock Price", "P/E (TTM)", "Forward P/E", "Price/Book",
        "Dividend Yield", "Market Cap ($B)", "52W High", "52W Low",
    ]

    table = doc.add_table(rows=len(metrics) + 1, cols=len(ticker_order) + 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    header_cells = ["Metric"] + ticker_order + ["Peer Median"]
    for i, h in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "003366")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for r_idx, metric in enumerate(metrics):
        table.rows[r_idx + 1].cells[0].text = metric
        for p in table.rows[r_idx + 1].cells[0].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.bold = True

        for c_idx, ticker in enumerate(ticker_order):
            val = tickers_data.get(ticker, {}).get(metric)
            if val is None:
                formatted = "[N/A]"
            elif metric in ("Stock Price", "52W High", "52W Low"):
                formatted = f"${val:.2f}"
            elif metric == "Dividend Yield":
                formatted = f"{val:.2f}%"
            elif metric == "Market Cap ($B)":
                formatted = f"{val:.1f}"
            else:
                formatted = f"{val:.2f}"

            cell = table.rows[r_idx + 1].cells[c_idx + 1]
            cell.text = formatted
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Peer median
        med_val = peer_median.get(metric)
        if med_val is None or med_val == "-":
            med_str = "-" if med_val == "-" else "[N/A]"
        elif metric in ("Stock Price", "52W High", "52W Low"):
            med_str = f"${med_val:.2f}"
        elif metric == "Dividend Yield":
            med_str = f"{med_val:.2f}%"
        elif metric == "Market Cap ($B)":
            med_str = f"{med_val:.1f}"
        else:
            med_str = f"{med_val:.2f}"

        cell = table.rows[r_idx + 1].cells[-1]
        cell.text = med_str
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def generate_word_doc(memo_data: dict) -> io.BytesIO:
    """Generate a formatted Word document from memo data.

    Returns a BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    analyst = memo_data.get("analyst", "Analyst")
    firm = memo_data.get("firm", "")
    date = memo_data.get("date", datetime.now().strftime("%B %d, %Y"))
    sections = memo_data.get("sections", {})

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"IR 1-Pager: {analyst} ({firm}) — Prepared {date}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    title.space_after = Pt(12)

    # Section 1: Bio
    if "bio" in sections:
        _add_section_header(doc, "Background and Analyst Bio")
        content = sections["bio"]
        if isinstance(content, dict):
            content = content.get("content", "")
        _parse_markdown_to_doc(doc, content)

    # Section 2: Forecast
    if "forecast" in sections:
        forecast = sections["forecast"]
        if isinstance(forecast, dict) and "table_rows" in forecast:
            date_updated = forecast.get("date_updated", "[Date unavailable]")
            _add_section_header(doc, f"{analyst}'s 3Q25E Forecast (As of {date_updated})")
            if forecast.get("is_stale"):
                p = doc.add_paragraph()
                run = p.add_run("Note: Forecast data is older than 30 days.")
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
            _add_forecast_table(doc, forecast)

    # Section 3: Earnings
    if "earnings" in sections:
        _add_section_header(doc, "Post-Earnings Feedback and Questions")
        content = sections["earnings"]
        if isinstance(content, dict):
            content = content.get("content", "")
        _parse_markdown_to_doc(doc, content)

    # Section 4: Peer Research
    if "peer" in sections:
        _add_section_header(doc, "Recent Peer Research")
        content = sections["peer"]
        if isinstance(content, dict):
            content = content.get("content", "")
        _parse_markdown_to_doc(doc, content)

    # Section 5: Valuation
    if "valuation" in sections:
        valuation = sections["valuation"]
        if isinstance(valuation, dict) and "tickers" in valuation:
            as_of = valuation.get("as_of", "")
            _add_section_header(doc, f"Valuation Ratios (As of {as_of})")
            _add_valuation_table(doc, valuation)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Confidential — For Internal Use Only")
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.name = "Calibri"

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
